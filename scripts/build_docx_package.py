#!/usr/bin/env python3
"""Build visually verified Word/PDF artifacts from Markdown sources."""

from __future__ import annotations

import argparse
from datetime import datetime, timezone
import hashlib
from importlib import metadata
import json
import os
import platform
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree
from pypdf import PdfReader, PdfWriter


BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(90, 90, 90)
ANI_ORANGE = RGBColor(255, 149, 0)
BLUE = RGBColor(46, 116, 181)
LIGHT_FILL = "F4F6F9"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"
DCTERMS_NS = "http://purl.org/dc/terms/"
MIN_ZIP_EPOCH = 315532800


def run(command: list[str], *, cwd: Path) -> None:
    subprocess.run(command, cwd=cwd, check=True)


def _source_date_epoch(value: int | None) -> int:
    raw = str(value) if value is not None else os.environ.get("SOURCE_DATE_EPOCH")
    if raw is None:
        raise SystemExit("Provide --source-date-epoch or set SOURCE_DATE_EPOCH")
    try:
        epoch = int(raw)
    except ValueError as exc:
        raise SystemExit("SOURCE_DATE_EPOCH must be an integer Unix timestamp") from exc
    if epoch < MIN_ZIP_EPOCH:
        raise SystemExit("SOURCE_DATE_EPOCH must be on or after 1980-01-01")
    return epoch


def _timestamp(epoch: int) -> str:
    return datetime.fromtimestamp(epoch, tz=timezone.utc).isoformat().replace("+00:00", "Z")


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _command_output(command: list[str], *, cwd: Path) -> str:
    return subprocess.check_output(command, cwd=cwd, text=True).strip()


def _command_version(executable: str, *arguments: str) -> str:
    output = subprocess.check_output(
        [executable, *arguments],
        text=True,
        stderr=subprocess.STDOUT,
    )
    return output.splitlines()[0].strip()


def _package_version(distribution: str) -> str:
    try:
        return metadata.version(distribution)
    except metadata.PackageNotFoundError:
        return "not_installed"


def tool_versions(*, pandoc: str, soffice: str) -> dict[str, str]:
    return {
        "python": platform.python_version(),
        "python_implementation": platform.python_implementation(),
        "pandoc": _command_version(pandoc, "--version"),
        "libreoffice": _command_version(soffice, "--version"),
        "python_docx": _package_version("python-docx"),
        "pypdf": _package_version("pypdf"),
        "lxml": _package_version("lxml"),
        "matplotlib": _package_version("matplotlib"),
        "numpy": _package_version("numpy"),
        "pillow": _package_version("Pillow"),
        "fonttools": _package_version("fonttools"),
    }


def git_source_state(root: Path) -> dict[str, object]:
    """Bind the receipt to the exact Git source state without claiming clean when dirty."""

    try:
        commit = _command_output(["git", "rev-parse", "HEAD"], cwd=root)
        tree = _command_output(["git", "rev-parse", "HEAD^{tree}"], cwd=root)
        status = subprocess.check_output(
            ["git", "status", "--porcelain=v1", "--untracked-files=all"],
            cwd=root,
        )
        tracked_patch = subprocess.check_output(
            ["git", "diff", "--binary", "--no-ext-diff", "--"],
            cwd=root,
        )
        staged_patch = subprocess.check_output(
            ["git", "diff", "--cached", "--binary", "--no-ext-diff", "--"],
            cwd=root,
        )
        untracked_raw = subprocess.check_output(
            ["git", "ls-files", "--others", "--exclude-standard", "-z"],
            cwd=root,
        )
    except (FileNotFoundError, subprocess.CalledProcessError) as exc:
        return {
            "available": False,
            "clean": None,
            "release_source_eligible": False,
            "blocker": f"git_source_state_unavailable:{type(exc).__name__}",
        }

    untracked = []
    for raw in untracked_raw.split(b"\0"):
        if not raw:
            continue
        relative = raw.decode("utf-8", errors="surrogateescape")
        candidate = root / relative
        untracked.append(
            {
                "path": relative,
                "sha256": _sha256(candidate) if candidate.is_file() else None,
            }
        )
    clean = not status
    return {
        "available": True,
        "commit": commit,
        "tree": tree,
        "clean": clean,
        "release_source_eligible": clean,
        "status_sha256": hashlib.sha256(status).hexdigest(),
        "status_lines": status.decode("utf-8", errors="replace").splitlines(),
        "tracked_worktree_patch_sha256": hashlib.sha256(tracked_patch).hexdigest(),
        "staged_patch_sha256": hashlib.sha256(staged_patch).hexdigest(),
        "untracked": untracked,
    }


def scrub_docx_package(path: Path, *, source_date_epoch: int | None = None) -> dict[str, int]:
    """Remove personal/session metadata before hashing or rendering the DOCX."""

    epoch = _source_date_epoch(source_date_epoch)
    timestamp = _timestamp(epoch)
    archive_datetime = datetime.fromtimestamp(epoch, tz=timezone.utc).timetuple()[:6]

    stats = {
        "rsid_attrs_removed": 0,
        "core_properties_scrubbed": 0,
        "custom_properties_removed": 0,
    }
    temporary = path.with_suffix(".privacy-scrub.tmp.docx")
    with zipfile.ZipFile(path, "r") as source:
        overrides: dict[str, bytes] = {}
        for name in source.namelist():
            story = bool(
                name == "word/document.xml"
                or re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)
                or name in {"word/footnotes.xml", "word/endnotes.xml"}
            )
            if not story:
                continue
            root = etree.fromstring(source.read(name))
            for element in root.iter():
                for attribute in list(element.attrib):
                    if attribute.startswith(f"{{{W_NS}}}rsid"):
                        del element.attrib[attribute]
                        stats["rsid_attrs_removed"] += 1
            overrides[name] = etree.tostring(
                root,
                xml_declaration=True,
                encoding="UTF-8",
                standalone=True,
            )

        if "docProps/core.xml" in source.namelist():
            root = etree.fromstring(source.read("docProps/core.xml"))
            changed = False
            for tag in (f"{{{DC_NS}}}creator", f"{{{CP_NS}}}lastModifiedBy"):
                for element in root.iter(tag):
                    if (element.text or "").strip():
                        element.text = ""
                        changed = True
            for local_name in ("created", "modified"):
                for element in root.iter(f"{{{DCTERMS_NS}}}{local_name}"):
                    if element.text != timestamp:
                        element.text = timestamp
                        changed = True
            if changed:
                stats["core_properties_scrubbed"] = 1
                overrides["docProps/core.xml"] = etree.tostring(
                    root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )

        if "_rels/.rels" in source.namelist():
            root = etree.fromstring(source.read("_rels/.rels"))
            changed = False
            for relationship in list(root.findall(f"{{{REL_NS}}}Relationship")):
                if (relationship.get("Target") or "").endswith("docProps/custom.xml"):
                    root.remove(relationship)
                    changed = True
            if changed:
                overrides["_rels/.rels"] = etree.tostring(
                    root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )

        if "[Content_Types].xml" in source.namelist():
            root = etree.fromstring(source.read("[Content_Types].xml"))
            changed = False
            for override in list(root.findall(f"{{{CT_NS}}}Override")):
                if (override.get("PartName") or "") == "/docProps/custom.xml":
                    root.remove(override)
                    changed = True
            if changed:
                overrides["[Content_Types].xml"] = etree.tostring(
                    root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )

        with zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as destination:
            for info in source.infolist():
                if info.filename == "docProps/custom.xml":
                    stats["custom_properties_removed"] = 1
                    continue
                canonical = zipfile.ZipInfo(info.filename, archive_datetime)
                canonical.compress_type = zipfile.ZIP_DEFLATED
                canonical.create_system = 3
                canonical.external_attr = info.external_attr
                canonical.internal_attr = info.internal_attr
                destination.writestr(
                    canonical,
                    overrides.get(info.filename, source.read(info.filename)),
                )
    temporary.replace(path)
    return stats


def font(run, *, name: str, size: float | None = None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(
    cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table) -> None:
    column_count = len(table.columns)
    width_maps = {
        1: [9360],
        2: [2500, 6860],
        3: [3000, 3180, 3180],
        4: [3600, 1920, 1920, 1920],
        5: [3000, 1590, 1590, 1590, 1590],
    }
    widths = width_maps.get(column_count)
    if widths is None:
        base = 9360 // column_count
        widths = [base] * column_count
        widths[-1] += 9360 - sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag, value in (("tblW", "9360"), ("tblInd", "120")):
        node = tbl_pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_pr.append(node)
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)
        cant_split.set(qn("w:val"), "true")
        if row_index == 0:
            header = tr_pr.find(qn("w:tblHeader"))
            if header is None:
                header = OxmlElement("w:tblHeader")
                tr_pr.append(header)
            header.set(qn("w:val"), "true")
        for column_index, cell in enumerate(row.cells):
            cell.width = int(widths[column_index] * 635)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is not None:
                tc_w.set(qn("w:w"), str(widths[column_index]))
                tc_w.set(qn("w:type"), "dxa")
            if row_index == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), LIGHT_FILL)
                cell._tc.get_or_add_tcPr().append(shd)


def style_docx(
    source: Path,
    destination: Path,
    *,
    preset: str,
    header_text: str,
    footer_text: str,
    figures: list[Path],
) -> int:
    doc = Document(source)
    is_publication = preset == "narrative_proposal"
    family = "Calibri"
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = family
    normal._element.rPr.rFonts.set(qn("w:ascii"), family)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), family)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8 if is_publication else 6)
    normal.paragraph_format.line_spacing = 1.333 if is_publication else 1.10
    # The publication preset deliberately uses a ragged-right body.  It avoids
    # the large inter-word gaps that Word/LibreOffice can create around inline
    # mathematics, URLs, code identifiers, and narrow table-adjacent paragraphs.
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    heading_tokens = {
        1: (16, 18 if is_publication else 16, 10 if is_publication else 8),
        2: (13, 12, 6),
        3: (12, 8, 4),
    }
    for level, (size, before, after) in heading_tokens.items():
        style = next(item for item in doc.styles if item.style_id == f"Heading{level}")
        style.font.name = family
        style._element.rPr.rFonts.set(qn("w:ascii"), family)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), family)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ANI_ORANGE if is_publication else BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for paragraph in doc.paragraphs:
        for item in paragraph.runs:
            font(item, name=family)
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True

    if doc.paragraphs:
        title = doc.paragraphs[0]
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(72 if is_publication else 20)
        title.paragraph_format.space_after = Pt(14)
        for item in title.runs:
            font(item, name=family, size=26 if is_publication else 22, color=BLACK, bold=True)
        metadata_end = 1
        if is_publication:
            metadata_end = next(
                (
                    index
                    for index, paragraph in enumerate(doc.paragraphs[1:], start=1)
                    if paragraph.style.name == "Heading 1"
                ),
                min(8, len(doc.paragraphs)),
            )
        else:
            metadata_end = min(2, len(doc.paragraphs))
        for paragraph in doc.paragraphs[1:metadata_end]:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(6)
            for item in paragraph.runs:
                font(item, name=family, size=10.5, color=MUTED, italic=True)
        if is_publication and metadata_end < len(doc.paragraphs):
            doc.paragraphs[metadata_end].paragraph_format.page_break_before = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run(header_text), name=family, size=8, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.add_run(f"{footer_text}  •  "), name=family, size=8, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    for table in doc.tables:
        table_geometry(table)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.05
                    for item in paragraph.runs:
                        font(item, name=family, size=8.5, bold=row_index == 0)

    if figures:
        doc.add_section(WD_SECTION.NEW_PAGE)
        heading = doc.add_paragraph("Figures", style="Heading 1")
        heading.paragraph_format.keep_with_next = True
        for index, figure in enumerate(figures, start=1):
            label = figure.stem.split("_", 1)[-1].replace("_", " ")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(figure), width=Inches(6.25))
            for doc_pr in run._element.xpath(".//wp:docPr"):
                doc_pr.set("title", f"AniBench Figure {index}")
                doc_pr.set("descr", f"AniBench chart: {label}")
            caption = doc.add_paragraph(f"Figure {index}. {label}")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.space_after = Pt(10)
            for item in caption.runs:
                font(item, name=family, size=9, color=MUTED, italic=True)
            if index < len(figures):
                doc.add_page_break()

    # Pandoc normally preserves Markdown alt text.  A deterministic fallback is
    # added only when an embedded image lacks one, so accessibility does not
    # depend on converter version.
    for index, shape in enumerate(doc.inline_shapes, start=1):
        doc_pr = shape._inline.docPr
        if not doc_pr.get("title"):
            doc_pr.set("title", f"AniBench figure {index}")
        if not doc_pr.get("descr"):
            kind = "method figure" if is_publication else "technical patent drawing aid"
            doc_pr.set(
                "descr",
                f"AniBench {kind} {index}; the adjacent manuscript text provides the full description.",
            )

    # Counsel-review drawing packets are easier to inspect when every drawing
    # starts on its own page.  Markdown images are already embedded by Pandoc;
    # set the containing paragraph rather than appending duplicate figures.
    if not is_publication and len(doc.inline_shapes) > 4:
        for paragraph in doc.paragraphs:
            if paragraph._p.xpath(".//w:drawing"):
                paragraph.paragraph_format.page_break_before = True
                paragraph.paragraph_format.keep_with_next = True

    # Authorship and inventorship are unresolved in the source documents; do not
    # infer either from repository ownership or document generation.
    doc.core_properties.author = ""
    doc.core_properties.keywords = (
        "AniBench, human trials, biological learning, evidence, acquisition"
    )
    embedded_figure_count = len(doc.inline_shapes)
    doc.save(destination)
    return embedded_figure_count


def normalize_pdf(
    source: Path,
    destination: Path,
    *,
    title: str,
    source_date_epoch: int,
) -> dict[str, object]:
    reader = PdfReader(source)
    writer = PdfWriter()
    # Clone the full document catalog so the PDF structure tree, marked-content
    # flags, bookmarks, language, and other accessibility objects survive the
    # deterministic metadata rewrite. Re-adding pages alone strips Tagged-PDF
    # semantics even when LibreOffice emitted them correctly.
    # Clone the complete catalog graph without copying LibreOffice's volatile
    # trailer /Info and /ID entries. The structure tree is rooted in the
    # catalog, so this preserves Tagged-PDF semantics while allowing the
    # normalized metadata and identifier below to be generated deterministically.
    writer.clone_reader_document_root(reader)
    timestamp = datetime.fromtimestamp(source_date_epoch, tz=timezone.utc).strftime(
        "D:%Y%m%d%H%M%SZ"
    )
    writer.metadata = {
        "/Title": title,
        "/Author": "",
        "/CreationDate": timestamp,
        "/ModDate": timestamp,
    }
    # LibreOffice writes a session timestamp into an XMP metadata stream and a
    # new trailer identifier on every conversion.  Clearing ``xmp_metadata``
    # removes the catalog reference, but ``clone_document_from_reader`` keeps
    # the now-unreferenced stream in the writer's object table unless it is
    # pruned explicitly.  Retaining either value makes byte-identical inputs
    # produce different PDFs even though the visible and tagged content match.
    writer.xmp_metadata = None
    writer.compress_identical_objects(
        remove_duplicates=False,
        remove_unreferenced=True,
    )
    writer.generate_file_identifiers()
    with destination.open("wb") as handle:
        writer.write(handle)
    normalized = PdfReader(destination)
    catalog = normalized.trailer["/Root"]
    mark_info = catalog.get("/MarkInfo") or {}
    if hasattr(mark_info, "get_object"):
        mark_info = mark_info.get_object()
    tagged_pdf = bool(mark_info.get("/Marked")) and catalog.get("/StructTreeRoot") is not None
    if not tagged_pdf:
        raise RuntimeError("normalized PDF lost its required Tagged-PDF structure")
    return {
        "page_count": len(normalized.pages),
        "pdf_sha256": hashlib.sha256(destination.read_bytes()).hexdigest(),
        "pdf_bytes": destination.stat().st_size,
        "tagged_pdf": True,
    }


def portable_receipt_path(path: Path, *, root: Path) -> str:
    """Keep public build receipts reproducible and free of workstation paths."""
    try:
        return path.relative_to(root).as_posix()
    except ValueError:
        return str(path)


def file_record(path: Path, *, root: Path) -> dict[str, object]:
    return {
        "path": portable_receipt_path(path, root=root),
        "sha256": _sha256(path),
        "bytes": path.stat().st_size,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--out-dir", type=Path, required=True)
    parser.add_argument("--stem", required=True)
    parser.add_argument("--title", required=True)
    parser.add_argument(
        "--preset", choices=("narrative_proposal", "standard_business_brief"), required=True
    )
    parser.add_argument("--header", required=True)
    parser.add_argument("--footer", required=True)
    parser.add_argument("--figures-dir", type=Path)
    parser.add_argument(
        "--figures-already-inline",
        action="store_true",
        help=(
            "hash and resolve --figures-dir assets without appending a duplicate "
            "figure plate after Markdown has already embedded them"
        ),
    )
    parser.add_argument("--source-date-epoch", type=int)
    args = parser.parse_args()

    root = Path(__file__).resolve().parents[1]
    source_date_epoch = _source_date_epoch(args.source_date_epoch)
    source = args.source.resolve()
    out = args.out_dir.resolve()
    out.mkdir(parents=True, exist_ok=True)
    figures = sorted(args.figures_dir.resolve().glob("*.png")) if args.figures_dir else []
    figures_to_append = [] if args.figures_already_inline else figures
    pandoc = shutil.which("pandoc")
    soffice = shutil.which("soffice")
    if not pandoc or not soffice:
        raise SystemExit("pandoc and soffice are required")

    docx_path = out / f"{args.stem}.docx"
    pdf_path = out / f"{args.stem}.pdf"
    with tempfile.TemporaryDirectory(prefix="anibench-docx-") as directory:
        temp = Path(directory)
        raw = temp / "raw.docx"
        run(
            [
                pandoc,
                str(source),
                "--from=markdown+tex_math_dollars+tex_math_single_backslash",
                "--to=docx",
                "--standalone",
                f"--resource-path={root}:{source.parent}:{args.figures_dir or source.parent}",
                "--output",
                str(raw),
            ],
            cwd=root,
        )
        embedded_figure_count = style_docx(
            raw,
            docx_path,
            preset=args.preset,
            header_text=args.header,
            footer_text=args.footer,
            figures=figures_to_append,
        )
        privacy_scrub = scrub_docx_package(
            docx_path,
            source_date_epoch=source_date_epoch,
        )
        office = temp / "office"
        office.mkdir()
        pdf_export_filter = (
            'pdf:writer_pdf_Export:{"UseTaggedPDF":{"type":"boolean","value":"true"}}'
        )
        run(
            [
                soffice,
                "--headless",
                "--convert-to",
                pdf_export_filter,
                "--outdir",
                str(office),
                str(docx_path),
            ],
            cwd=root,
        )
        rendered = office / f"{args.stem}.pdf"
        if not rendered.exists():
            raise RuntimeError("LibreOffice did not create the expected PDF")
        receipt = normalize_pdf(
            rendered,
            pdf_path,
            title=args.title,
            source_date_epoch=source_date_epoch,
        )

    builder_paths = [Path(__file__).resolve()]
    method_figure_builder = root / "paper" / "v2" / "build_method_figures.py"
    if args.figures_dir and method_figure_builder.is_file():
        builder_paths.append(method_figure_builder)
    figure_assets = (
        sorted(
            path
            for path in args.figures_dir.resolve().iterdir()
            if path.is_file() and path.suffix.lower() in {".png", ".svg"}
        )
        if args.figures_dir
        else []
    )

    receipt.update(
        {
            "schema_version": "anibench.document-build-receipt.v2",
            "source_date_epoch": source_date_epoch,
            "build_timestamp_utc": _timestamp(source_date_epoch),
            "preset": args.preset,
            "build_parameters": {
                "stem": args.stem,
                "title": args.title,
                "preset": args.preset,
                "header": args.header,
                "footer": args.footer,
                "figures_dir": (
                    portable_receipt_path(args.figures_dir.resolve(), root=root)
                    if args.figures_dir
                    else None
                ),
                "figures_already_inline": args.figures_already_inline,
                "pandoc_from": "markdown+tex_math_dollars+tex_math_single_backslash",
                "pandoc_to": "docx",
                "libreoffice_conversion": "headless_docx_to_pdf",
                "pdf_export_filter_data": {
                    "filter": "writer_pdf_Export",
                    "UseTaggedPDF": True,
                },
            },
            "layout_override": (
                "publication_readability_left_align"
                if args.preset == "narrative_proposal"
                else "counsel_drawings_one_per_page_when_embedded_count_gt_4"
            ),
            "source": portable_receipt_path(source, root=root),
            "source_sha256": hashlib.sha256(source.read_bytes()).hexdigest(),
            "docx": portable_receipt_path(docx_path, root=root),
            "docx_sha256": hashlib.sha256(docx_path.read_bytes()).hexdigest(),
            "pdf": portable_receipt_path(pdf_path, root=root),
            "figure_count": embedded_figure_count,
            "appended_figure_count": len(figures_to_append),
            "figure_assets": [file_record(path, root=root) for path in figure_assets],
            "builders": [file_record(path, root=root) for path in builder_paths],
            "tool_versions": tool_versions(pandoc=pandoc, soffice=soffice),
            "git_source_state": git_source_state(root),
            "privacy_scrub": privacy_scrub,
        }
    )
    receipt_path = out / f"{args.stem}_build_receipt.json"
    receipt_path.write_text(json.dumps(receipt, indent=2, sort_keys=True) + "\n", encoding="utf-8")
    print(json.dumps(receipt, indent=2, sort_keys=True))


if __name__ == "__main__":
    main()
