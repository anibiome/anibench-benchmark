from __future__ import annotations

from datetime import datetime, timezone
import hashlib
import json
import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document
from lxml import etree
from pypdf import PdfReader

from scripts.build_docx_package import scrub_docx_package, table_geometry


ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "paper" / "v2" / "AniBench_v2_benchmark_protocol.md"
CHECKLIST = ROOT / "paper" / "v2" / "REPORTING_CHECKLIST.md"
LEVEL1_REQUIREMENTS = (
    ROOT / "spec" / "v3" / "level1" / "role-aware-target-requirements.v3.json"
)
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"
DCTERMS_NS = "http://purl.org/dc/terms/"
SOURCE_DATE_EPOCH = 1_783_641_600


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def test_docx_build_privacy_scrub_removes_personal_and_custom_metadata(
    tmp_path: Path,
) -> None:
    path = tmp_path / "candidate.docx"
    document = Document()
    document.add_heading("AniBench", level=1)
    document.core_properties.author = "Private Author"
    document.core_properties.last_modified_by = "Private Editor"
    document.save(path)
    with zipfile.ZipFile(path, "a", zipfile.ZIP_DEFLATED) as package:
        package.writestr(
            "docProps/custom.xml",
            (
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Properties xmlns="http://schemas.openxmlformats.org/'
                'officeDocument/2006/custom-properties"/>'
            ),
        )

    receipt = scrub_docx_package(path, source_date_epoch=SOURCE_DATE_EPOCH)
    assert receipt["custom_properties_removed"] == 1
    assert receipt["core_properties_scrubbed"] == 1
    with zipfile.ZipFile(path) as package:
        assert "docProps/custom.xml" not in package.namelist()
        core = etree.fromstring(package.read("docProps/core.xml"))
        assert (core.findtext(f"{{{DC_NS}}}creator") or "") == ""
        assert (core.findtext(f"{{{CP_NS}}}lastModifiedBy") or "") == ""
        timestamp = datetime.fromtimestamp(SOURCE_DATE_EPOCH, tz=timezone.utc).isoformat().replace(
            "+00:00", "Z"
        )
        assert core.findtext(f"{{{DCTERMS_NS}}}created") == timestamp
        assert core.findtext(f"{{{DCTERMS_NS}}}modified") == timestamp
        document_xml = etree.fromstring(package.read("word/document.xml"))
        assert not any(
            attribute.startswith(f"{{{W_NS}}}rsid")
            for element in document_xml.iter()
            for attribute in element.attrib
        )
        expected_archive_time = datetime.fromtimestamp(
            SOURCE_DATE_EPOCH,
            tz=timezone.utc,
        ).timetuple()[:6]
        assert {item.date_time for item in package.infolist()} == {expected_archive_time}


def test_docx_tables_prevent_rows_from_splitting_across_pages() -> None:
    document = Document()
    table = document.add_table(rows=3, cols=3)

    table_geometry(table)

    for row in table.rows:
        cant_split = row._tr.get_or_add_trPr().find(f"{{{W_NS}}}cantSplit")
        assert cant_split is not None
        assert cant_split.get(f"{{{W_NS}}}val") == "true"


def test_method_figures_are_two_build_deterministic_and_epoch_bound(
    tmp_path: Path,
) -> None:
    output = tmp_path / "figures"
    command = [
        sys.executable,
        str(ROOT / "paper" / "v2" / "build_method_figures.py"),
        "--out-dir",
        str(output),
        "--source-date-epoch",
        str(SOURCE_DATE_EPOCH),
    ]
    subprocess.run(command, cwd=ROOT, check=True)
    first = {path.name: _sha256(path) for path in sorted(output.iterdir())}
    subprocess.run(command, cwd=ROOT, check=True)
    second = {path.name: _sha256(path) for path in sorted(output.iterdir())}

    assert first == second
    assert len(first) == 8
    assert {Path(name).suffix for name in first} == {".png", ".svg"}
    timestamp = datetime.fromtimestamp(SOURCE_DATE_EPOCH, tz=timezone.utc).isoformat().replace(
        "+00:00", "Z"
    )
    for path in output.iterdir():
        assert int(path.stat().st_mtime) == SOURCE_DATE_EPOCH
    for svg in output.glob("*.svg"):
        text = svg.read_text(encoding="utf-8")
        assert timestamp in text
        assert "anibench-v2-method-figures" not in text
        assert "AniBench v2 deterministic method-figure builder" in text


@pytest.mark.skipif(
    shutil.which("pandoc") is None or shutil.which("soffice") is None,
    reason="document parity requires pandoc and LibreOffice",
)
def test_document_build_is_two_pass_deterministic_with_complete_provenance(
    tmp_path: Path,
) -> None:
    figures = tmp_path / "figures"
    subprocess.run(
        [
            sys.executable,
            str(ROOT / "paper" / "v2" / "build_method_figures.py"),
            "--out-dir",
            str(figures),
            "--source-date-epoch",
            str(SOURCE_DATE_EPOCH),
        ],
        cwd=ROOT,
        check=True,
    )
    source = tmp_path / "source.md"
    source.write_text(
        "# AniBench deterministic build fixture\n\n"
        "## Methods\n\nThe receipt must bind every builder and method figure.\n",
        encoding="utf-8",
    )
    output = tmp_path / "build"
    command = [
        sys.executable,
        str(ROOT / "scripts" / "build_docx_package.py"),
        "--source",
        str(source),
        "--out-dir",
        str(output),
        "--stem",
        "fixture",
        "--title",
        "AniBench deterministic build fixture",
        "--preset",
        "narrative_proposal",
        "--header",
        "ANIBENCH / TEST",
        "--footer",
        "ANIBENCH / TEST",
        "--figures-dir",
        str(figures),
        "--source-date-epoch",
        str(SOURCE_DATE_EPOCH),
    ]

    subprocess.run(command, cwd=ROOT, check=True, capture_output=True, text=True)
    receipt_path = output / "fixture_build_receipt.json"
    first_receipt = json.loads(receipt_path.read_text(encoding="utf-8"))
    first_docx = _sha256(output / "fixture.docx")
    first_pdf = _sha256(output / "fixture.pdf")

    subprocess.run(command, cwd=ROOT, check=True, capture_output=True, text=True)
    second_receipt = json.loads(receipt_path.read_text(encoding="utf-8"))

    assert _sha256(output / "fixture.docx") == first_docx
    assert _sha256(output / "fixture.pdf") == first_pdf
    assert second_receipt == first_receipt
    assert first_receipt["schema_version"] == "anibench.document-build-receipt.v2"
    assert first_receipt["source_date_epoch"] == SOURCE_DATE_EPOCH
    assert first_receipt["build_timestamp_utc"].endswith("Z")
    assert first_receipt["build_parameters"] == {
        "stem": "fixture",
        "title": "AniBench deterministic build fixture",
        "preset": "narrative_proposal",
        "header": "ANIBENCH / TEST",
        "footer": "ANIBENCH / TEST",
        "figures_dir": str(figures),
        "figures_already_inline": False,
        "pandoc_from": "markdown+tex_math_dollars+tex_math_single_backslash",
        "pandoc_to": "docx",
        "libreoffice_conversion": "headless_docx_to_pdf",
        "pdf_export_filter_data": {
            "filter": "writer_pdf_Export",
            "UseTaggedPDF": True,
        },
    }
    assert first_receipt["docx_sha256"] == first_docx
    assert first_receipt["pdf_sha256"] == first_pdf
    assert len(first_receipt["figure_assets"]) == 8
    assert len(first_receipt["builders"]) == 2
    assert {Path(row["path"]).name for row in first_receipt["builders"]} == {
        "build_docx_package.py",
        "build_method_figures.py",
    }
    assert all(len(row["sha256"]) == 64 for row in first_receipt["figure_assets"])
    assert all(len(row["sha256"]) == 64 for row in first_receipt["builders"])
    for row in first_receipt["figure_assets"]:
        assert row["sha256"] == _sha256(Path(row["path"]))
        assert row["bytes"] == Path(row["path"]).stat().st_size
    for row in first_receipt["builders"]:
        builder = ROOT / row["path"]
        assert row["sha256"] == _sha256(builder)
        assert row["bytes"] == builder.stat().st_size
    assert set(first_receipt["tool_versions"]) == {
        "python",
        "python_implementation",
        "pandoc",
        "libreoffice",
        "python_docx",
        "pypdf",
        "lxml",
        "matplotlib",
        "numpy",
        "pillow",
        "fonttools",
    }
    assert "not_installed" not in first_receipt["tool_versions"].values()
    pdf_metadata = PdfReader(output / "fixture.pdf").metadata
    expected_pdf_date = datetime.fromtimestamp(
        SOURCE_DATE_EPOCH,
        tz=timezone.utc,
    ).strftime("D:%Y%m%d%H%M%SZ")
    assert pdf_metadata.creation_date == datetime.fromtimestamp(
        SOURCE_DATE_EPOCH,
        tz=timezone.utc,
    )
    assert pdf_metadata.get("/CreationDate") == expected_pdf_date
    assert pdf_metadata.get("/ModDate") == expected_pdf_date
    source_state = first_receipt["git_source_state"]
    assert source_state["available"] is True
    assert len(source_state["commit"]) == 40
    assert len(source_state["tree"]) == 40
    assert len(source_state["status_sha256"]) == 64
    assert len(source_state["tracked_worktree_patch_sha256"]) == 64
    assert len(source_state["staged_patch_sha256"]) == 64
    assert all(
        row["sha256"] is None or len(row["sha256"]) == 64
        for row in source_state["untracked"]
    )
    assert source_state["release_source_eligible"] is source_state["clean"]


def test_make_paper_regenerates_and_binds_method_figures() -> None:
    makefile = _text(ROOT / "Makefile")
    paper_recipe = makefile.split("paper:\n", 1)[1].split("\nserve-studio:", 1)[0]
    assert "paper/v2/build_method_figures.py" in paper_recipe
    assert "--figures-dir paper/v2/figures" in paper_recipe
    assert paper_recipe.index("build_method_figures.py") < paper_recipe.index(
        "build_docx_package.py"
    )
    assert paper_recipe.count('--source-date-epoch "$(SOURCE_DATE_EPOCH)"') == 2


def test_v2_protocol_contains_required_methods_sections() -> None:
    text = _text(PAPER)
    required = (
        "## 3. Six benchmark families",
        "## 5. Participant-event information",
        "## 7. Causal, sequential, and transport geometry",
        "## 8. Typed uncertainty and evidence lanes",
        "## 9. Level-1 authority",
        "## 10. Source-bound protocol compilation",
        "## 11. Prospective protocol optimization",
        "## 13. Validation protocol",
        "### 13.4 Prospective hypotheses awaiting immutable registration",
        "## 14. Anti-gaming and mathematical invariants",
        "## 16. Governance and promotion gates",
        "## 18. Reproducibility map",
        "## 19. Limitations and next empirical step",
        "## 20. Declarations and submission gates",
        "## References",
    )
    for heading in required:
        assert heading in text


def test_v2_protocol_binds_every_formula_and_contract_path() -> None:
    text = _text(PAPER)
    required_paths = (
        "spec/v2/level1/biological-coordinate-registry.json",
        "spec/v3/level1/role-aware-target-requirements.v3.json",
        "spec/v3/level1/migrations/v2-to-v3-substantive-impact-receipt.json",
        "schemas/v2/protocol-capacity-input.schema.json",
        "schemas/v2/optimizer-protocol-input.schema.json",
        "src/anibench/information_v2.py",
        "src/anibench/causal_v2.py",
        "src/anibench/protocol_capacity_v2.py",
        "src/anibench/optimizer_protocol_v2.py",
        "src/anibench/source_atlas_v2.py",
        "data/source_projections/v2/EXTERNAL_SOURCE_ACQUISITION_LEDGER.json",
        "packaging/public_v2/EXTERNAL_SOURCE_VERIFICATION_RECEIPT.json",
        "packaging/public_v2/EXTERNAL_FIELD_PROVENANCE_RECEIPT.json",
        "scripts/verify_external_field_receipts.py",
        "src/anibench/api.py",
        "src/anibench/cli.py",
        "src/anibench/studio.py",
    )
    for relative in required_paths:
        assert relative in text
        assert (ROOT / relative).is_file()


@pytest.mark.parametrize("path", (PAPER, CHECKLIST))
def test_v2_publication_candidate_contains_no_current_rank_numbers_or_claims(
    path: Path,
) -> None:
    text = _text(path)
    forbidden_current_values = (
        "54.870738",
        "34.080014",
        "46.991512",
        "15.908317",
        "20.598901",
        "0.014074",
        "0.002393",
    )
    for value in forbidden_current_values:
        assert value not in text

    forbidden_claim_patterns = (
        r"\brank(?:ed)?\s*(?:#|number\s*)?\d+\b",
        r"\b(?:first|second|third)[ -](?:ranked|place)\b",
        r"\bworld(?:wide)?\s+rank\b",
        r"\bpublic[ -]ready\b",
        r"\bready\s+for\s+(?:public|publication|release)\b",
        r"\b(?:is|are|has been|have been)\s+(?:empirically\s+)?validated\b",
        r"\bwe\s+(?:show|demonstrate|find|report)\s+that\b",
        r"\boutperforms?\b",
        r"^#{1,6}\s+authors?\b",
        r"^authors?\s*:",
        r"\bdoi\s*:",
        r"\bdoi\s+(?:assigned|registered)\b",
        r"\bstate-of-the-art\s+benchmark\b",
    )
    for pattern in forbidden_claim_patterns:
        assert re.search(pattern, text, flags=re.IGNORECASE | re.MULTILINE) is None


def test_v2_protocol_uses_protocol_language_for_empirical_hypotheses() -> None:
    text = _text(PAPER)
    assert "The following are hypotheses, not findings." in text
    assert "Empirical validation state: `not_run`" in text
    assert "Trial-ranking results: absent" in text
    assert "Public release decision: absent" in text
    assert "not yet registered" in text
    assert "[AUTHOR LIST AND ORDER REQUIRE CONTRIBUTOR CONFIRMATION]" in text
    assert "[AFFILIATIONS REQUIRE AUTHOR CONFIRMATION]" in text
    assert "[CORRESPONDING AUTHOR AND CONTACT REQUIRE CONFIRMATION]" in text


def test_v2_protocol_matches_registered_stage_set_and_decision_epoch_contract() -> None:
    text = _text(PAPER)
    prose = " ".join(text.split())
    assert "one explicitly\nregistered stage set" in text
    assert "An unbound stage remains a singleton alternative and is never\nadded" in text
    assert "sole supported aggregation rule sums complete policy" in text
    assert "declared mutually disjoint" in text
    assert "One stage cannot appear in two registered sets" in text
    assert "never constructs a synthetic best case" in text
    assert "requires either an explicit decision-epoch ledger or" in text
    assert "registered regular decision-epoch process" in text
    assert "produce equivalent support geometry" in text
    assert "\\widetilde\\omega_{ek}=\\omega_{ek}\\frac{n_e^M}{n_e^Y}" in text
    assert "\\text{SMART epoch-specific estimands}" in text
    assert "\\text{pooled micro-randomized estimand}" in text
    assert "exactly equivalent to its explicit\ndecision grid" in text
    assert "It is not biological information or inferential\nprecision" in text
    assert "state-to-policy contrast operator" in text
    assert "Neither gate estimates a moderation effect" in text
    assert "column-orthonormal" in text
    assert "diagonalize the frozen prior-whitened" in text
    assert "registered crosswalk remains unresolved for numeric credit" in prose
    assert "does not estimate a transported effect" in prose
    assert "Every transport-axis family binds" in prose
    assert "Extra declared coordinates are ignored numerically" in prose
    assert "with multiple families they are null" in prose
    assert "No family is winner-selected" in prose

    forbidden = (
        "Stages are aggregated",
        "coherent stage-class aggregation",
        "within-class aggregator",
        "only coherent stage classes are aggregated",
        "stage-level surrogate",
        "does not yet instantiate",
    )
    for phrase in forbidden:
        assert phrase not in text


def test_v2_reporting_checklist_matches_registered_stage_set_contract() -> None:
    text = _text(CHECKLIST)
    prose = " ".join(text.split())
    assert "authority-registered stage set" in prose
    assert "participant sets are mutually disjoint" in prose
    assert "One stage cannot belong to multiple aggregation sets" in prose
    assert "sums its complete aligned policy, component, and moderator" in prose
    assert "no cross-stage independence/dependence object" not in prose
    assert "perform no cross-stage" not in prose


def test_v2_protocol_matches_role_aware_level1_authority() -> None:
    text = _text(PAPER)
    prose = " ".join(text.split())
    for required in (
        "One scientific map, seven estimation roles",
        "Direct mutable outcome basis",
        "`D`, `P`, `H`, and `T` coordinates cannot be coerced",
        "Six family-specific operating-characteristic authorities",
        "Level 1 does not declare a global participant target",
        "source-bound joint\ncontext-support object",
        "v2-to-v3 substantive-impact receipt",
    ):
        assert required in text
    for superseded in (
        "193{,}536",
        "1{,}544{,}148",
        "171,572 participants",
        "level1_target_percent",
        "level1_uncapped_ratio",
    ):
        assert superseded not in text
    assert "family-specific operating-characteristic and enrollment derivations" in prose
    assert "global enrollment value may be emitted only after all six" in prose


def test_v2_protocol_level1_roles_match_normative_target_source_object() -> None:
    text = _text(PAPER)
    source = json.loads(_text(LEVEL1_REQUIREMENTS))
    roles = source["estimation_role_authority"]["roles"]
    membership = [coordinate for role in roles for coordinate in role["coordinate_ids"]]
    assert len(roles) == 7
    assert len(membership) == 64
    assert len(set(membership)) == 64
    direct = next(role for role in roles if role["role_id"] == "direct_mutable_outcome_basis")
    assert len(direct["coordinate_ids"]) == 22
    assert source["enrollment_authority"]["global_enrollment"]["state"] == "unresolved"
    assert source["enrollment_authority"]["global_enrollment"]["value"] is None
    assert source["noncompensatory_family_vector"]["aggregation"]["state"] == "forbidden"
    assert source["noncompensatory_family_vector"]["rank"]["state"] == "forbidden"
    assert "64-coordinate scientific" in text
    assert "seven disjoint estimation" in text


def test_v2_protocol_matches_typed_coordinate_and_public_export_boundaries() -> None:
    text = _text(PAPER)
    assert "generic nonnegative coordinate contract has three numerical forms" in text
    assert "Unknown or absent geometry is not encoded as a fabricated numeric coordinate" in text
    assert "affected family outputs are null or unresolved rather than zero" in text
    assert "fresh-history public export contains a score-free external source atlas" in text
    assert "Controlled ANI projections" in text
    assert "private trial data" in text
    assert "are excluded\nfrom the public repository and distribution archives" in text
    assert "The initial source atlas contains ANI-affiliated studies" not in text


def test_reporting_checklist_covers_all_required_evidence_classes() -> None:
    text = _text(CHECKLIST)
    for heading in (
        "## D. Participant-event manifest",
        "## E. Intervention and estimand geometry",
        "## F. Joint information computation",
        "## I. Typed uncertainty",
        "## J. Evidence lanes and access",
        "## K. Role-aware Level-1 authority",
        "## L. Validation and preregistration",
        "## M. Anti-gaming and numerical invariants",
        "## O. Governance, independence, and release",
    ):
        assert heading in text


def test_protocol_references_primary_methods_sources_already_routed_by_repo() -> None:
    text = _text(PAPER)
    primary_sources = (
        "https://doi.org/10.1214/aoms/1177728069",
        "https://doi.org/10.1214/ss/1177009939",
        "https://doi.org/10.4153/CJM-1960-030-4",
        "https://database.ich.org/sites/default/files/E9-R1_Step4_Guideline_2019_1203.pdf",
        "https://www.fda.gov/regulatory-information/search-fda-guidance-documents/adaptive-design-clinical-trials-drugs-and-biologics-guidance-industry",
        "https://doi.org/10.1093/aje/kwq084",
        "https://doi.org/10.1038/sdata.2016.18",
        "https://doi.org/10.6028/NIST.AI.800-2.ipd",
        "https://doi.org/10.6028/NIST.AI.800-3",
        "https://doi.org/10.1037/hea0000305",
        "https://doi.org/10.1093/biomet/asaa070",
        "https://doi.org/10.1002/sim.9748",
        "https://doi.org/10.1093/biostatistics/kxaf003",
    )
    for source in primary_sources:
        assert source in text


def test_publication_docs_use_the_role_aware_level1_contract() -> None:
    paths = (
        PAPER,
        CHECKLIST,
        ROOT / "README.md",
        ROOT / "CURRENT.md",
        ROOT / "docs" / "API.md",
        ROOT / "docs" / "DESIGN_STUDIO.md",
        ROOT / "docs" / "RELEASE_CHECKLIST.md",
    )
    texts = {path: _text(path) for path in paths}
    combined = "\n".join(texts.values())

    assert "1,544,148" not in combined
    assert "1544148" not in combined
    assert "193,536" not in combined
    assert "193536" not in combined
    assert "Protocol-capacity input and compiler" in texts[PAPER]
    assert "role-aware Level-1" in combined
    assert "seven disjoint" in combined
    assert "22 direct" in combined or "twenty-two direct" in combined
    assert "level1_target_percent" not in texts[PAPER]
    assert "level1_uncapped_ratio" not in texts[PAPER]
    assert "maturity penalty" in texts[ROOT / "docs" / "DESIGN_STUDIO.md"]
    assert "future\n> superintelligence" in texts[ROOT / "README.md"]
    assert "recover a usable model of latent biological state" in texts[ROOT / "README.md"]
    assert "overall scalar or rank" in texts[ROOT / "README.md"]
